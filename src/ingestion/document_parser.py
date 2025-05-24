"""
Parser de Documentos - Suporte para múltiplos formatos
Extrai texto de diferentes tipos de documentos.
"""

import io
import logging
from pathlib import Path
from typing import Dict, List, Optional, Any
from abc import ABC, abstractmethod

# Importações condicionais para diferentes formatos
try:
    import PyPDF2
    from pdfplumber import PDF
    HAS_PDF_SUPPORT = True
except ImportError:
    HAS_PDF_SUPPORT = False

try:
    from docx import Document
    HAS_DOCX_SUPPORT = True
except ImportError:
    HAS_DOCX_SUPPORT = False

try:
    import markdown
    HAS_MARKDOWN_SUPPORT = True
except ImportError:
    HAS_MARKDOWN_SUPPORT = False

logger = logging.getLogger(__name__)

class BaseParser(ABC):
    """Classe base para parsers de documentos."""
    
    @abstractmethod
    def parse(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Parse o conteúdo do arquivo e retorna texto e metadados.
        
        Args:
            file_content: Conteúdo do arquivo em bytes
            filename: Nome do arquivo
            
        Returns:
            Dict com 'text', 'metadata' e 'pages' (se aplicável)
        """
        pass
    
    @abstractmethod
    def supports_format(self, filename: str) -> bool:
        """Verifica se o parser suporta o formato do arquivo."""
        pass

class TextParser(BaseParser):
    """Parser para arquivos de texto simples (.txt)."""
    
    def parse(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Parse arquivos de texto simples."""
        try:
            # Tenta diferentes encodings
            encodings = ['utf-8', 'latin1', 'cp1252']
            text = None
            encoding_used = None
            
            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    encoding_used = encoding
                    break
                except UnicodeDecodeError:
                    continue
            
            if text is None:
                raise ValueError("Não foi possível decodificar o arquivo de texto")
            
            return {
                'text': text.strip(),
                'metadata': {
                    'encoding': encoding_used,
                    'char_count': len(text),
                    'line_count': len(text.splitlines())
                },
                'pages': 1
            }
            
        except Exception as e:
            logger.error(f"Erro ao processar arquivo TXT {filename}: {str(e)}")
            raise

    def supports_format(self, filename: str) -> bool:
        """Suporta arquivos .txt."""
        return filename.lower().endswith('.txt')

class MarkdownParser(BaseParser):
    """Parser para arquivos Markdown (.md)."""
    
    def parse(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Parse arquivos Markdown."""
        try:
            # Decodifica o conteúdo
            markdown_text = file_content.decode('utf-8')
            
            result = {
                'text': markdown_text.strip(),
                'metadata': {
                    'format': 'markdown',
                    'char_count': len(markdown_text),
                    'line_count': len(markdown_text.splitlines())
                },
                'pages': 1
            }
            
            # Se markdown está disponível, converte para HTML também
            if HAS_MARKDOWN_SUPPORT:
                md = markdown.Markdown(extensions=['meta', 'toc'])
                html_content = md.convert(markdown_text)
                
                result['metadata'].update({
                    'html_content': html_content,
                    'meta': md.Meta if hasattr(md, 'Meta') else {}
                })
            
            return result
            
        except Exception as e:
            logger.error(f"Erro ao processar arquivo MD {filename}: {str(e)}")
            raise

    def supports_format(self, filename: str) -> bool:
        """Suporta arquivos .md."""
        return filename.lower().endswith('.md')

class PDFParser(BaseParser):
    """Parser para arquivos PDF."""
    
    def parse(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Parse arquivos PDF."""
        if not HAS_PDF_SUPPORT:
            raise ImportError("Bibliotecas de PDF não estão instaladas")
        
        try:
            pdf_file = io.BytesIO(file_content)
            text_content = []
            metadata = {}
            
            # Tenta primeiro com pdfplumber (melhor para texto)
            try:
                with PDF(pdf_file) as pdf:
                    for page_num, page in enumerate(pdf.pages):
                        page_text = page.extract_text()
                        if page_text:
                            text_content.append(f"[Página {page_num + 1}]\n{page_text}")
                    
                    metadata.update({
                        'page_count': len(pdf.pages),
                        'parser_used': 'pdfplumber'
                    })
                    
            except Exception as pdfplumber_error:
                logger.warning(f"pdfplumber falhou, tentando PyPDF2: {pdfplumber_error}")
                
                # Fallback para PyPDF2
                pdf_file.seek(0)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(f"[Página {page_num + 1}]\n{page_text}")
                
                metadata.update({
                    'page_count': len(pdf_reader.pages),
                    'parser_used': 'PyPDF2',
                    'metadata': dict(pdf_reader.metadata) if pdf_reader.metadata else {}
                })
            
            full_text = "\n\n".join(text_content)
            
            if not full_text.strip():
                raise ValueError("Não foi possível extrair texto do PDF")
            
            return {
                'text': full_text.strip(),
                'metadata': metadata,
                'pages': metadata.get('page_count', 0)
            }
            
        except Exception as e:
            logger.error(f"Erro ao processar arquivo PDF {filename}: {str(e)}")
            raise

    def supports_format(self, filename: str) -> bool:
        """Suporta arquivos .pdf."""
        return filename.lower().endswith('.pdf')

class DOCXParser(BaseParser):
    """Parser para arquivos Word (.docx)."""
    
    def parse(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Parse arquivos DOCX."""
        if not HAS_DOCX_SUPPORT:
            raise ImportError("python-docx não está instalado")
        
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            
            # Extrai texto dos parágrafos
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text.strip())
            
            # Extrai texto das tabelas
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_texts.append(cell.text.strip())
                    if row_texts:
                        table_texts.append(" | ".join(row_texts))
            
            # Combina todo o texto
            all_text = "\n\n".join(paragraphs)
            if table_texts:
                all_text += "\n\n[TABELAS]\n" + "\n".join(table_texts)
            
            # Extrai propriedades do documento
            core_props = doc.core_properties
            metadata = {
                'paragraph_count': len(paragraphs),
                'table_count': len(doc.tables),
                'author': core_props.author if core_props.author else None,
                'title': core_props.title if core_props.title else None,
                'subject': core_props.subject if core_props.subject else None,
                'created': core_props.created.isoformat() if core_props.created else None,
                'modified': core_props.modified.isoformat() if core_props.modified else None
            }
            
            return {
                'text': all_text.strip(),
                'metadata': metadata,
                'pages': 1  # DOCX não tem conceito de páginas bem definido
            }
            
        except Exception as e:
            logger.error(f"Erro ao processar arquivo DOCX {filename}: {str(e)}")
            raise

    def supports_format(self, filename: str) -> bool:
        """Suporta arquivos .docx."""
        return filename.lower().endswith('.docx')

class DocumentParser:
    """Classe principal para parsing de documentos."""
    
    def __init__(self):
        """Inicializa o parser com todos os formatos suportados."""
        self.parsers: List[BaseParser] = [
            TextParser(),
            MarkdownParser(),
            PDFParser(),
            DOCXParser()
        ]
        
        logger.info("DocumentParser inicializado com os seguintes formatos:")
        for parser in self.parsers:
            logger.info(f"  - {parser.__class__.__name__}")
    
    def get_supported_formats(self) -> List[str]:
        """Retorna lista de formatos suportados."""
        formats = []
        test_files = ['test.txt', 'test.md', 'test.pdf', 'test.docx']
        
        for test_file in test_files:
            for parser in self.parsers:
                if parser.supports_format(test_file):
                    extension = Path(test_file).suffix
                    if extension not in formats:
                        formats.append(extension)
        
        return formats
    
    def parse_document(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Parse um documento usando o parser apropriado.
        
        Args:
            file_content: Conteúdo do arquivo em bytes
            filename: Nome do arquivo
            
        Returns:
            Dict com texto extraído e metadados
            
        Raises:
            ValueError: Se o formato não for suportado
            Exception: Erros durante o parsing
        """
        # Encontra o parser apropriado
        suitable_parser = None
        for parser in self.parsers:
            if parser.supports_format(filename):
                suitable_parser = parser
                break
        
        if suitable_parser is None:
            supported = self.get_supported_formats()
            raise ValueError(
                f"Formato não suportado para {filename}. "
                f"Formatos suportados: {', '.join(supported)}"
            )
        
        logger.info(f"Usando {suitable_parser.__class__.__name__} para {filename}")
        
        try:
            result = suitable_parser.parse(file_content, filename)
            
            # Adiciona metadados gerais
            result['metadata'].update({
                'filename': filename,
                'file_size': len(file_content),
                'parser_class': suitable_parser.__class__.__name__
            })
            
            logger.info(f"Documento {filename} processado com sucesso")
            return result
            
        except Exception as e:
            logger.error(f"Falha ao processar {filename}: {str(e)}")
            raise
    
    def validate_file(self, filename: str, max_size: int = 50 * 1024 * 1024) -> bool:
        """
        Valida se um arquivo pode ser processado.
        
        Args:
            filename: Nome do arquivo
            max_size: Tamanho máximo em bytes (padrão: 50MB)
            
        Returns:
            True se o arquivo pode ser processado
        """
        # Verifica formato
        if not any(parser.supports_format(filename) for parser in self.parsers):
            return False
        
        return True 